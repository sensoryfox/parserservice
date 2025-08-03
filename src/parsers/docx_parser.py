from __future__ import annotations

import re
from io import BytesIO
from uuid import UUID, uuid4
from typing import List

from docx import Document  # python-docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore

from ..models import Line, ImageArtefact, ParseResult
from .base import BaseParser


_HEADING_RE = re.compile(r"heading\s*([0-9]+)", re.I)


class DocxParser(BaseParser):
    """Парсер DOCX-файлов (абзацы, заголовки, таблицы, картинки)."""

    async def parse(
        self,
        *,
        doc_id: UUID,
        file_content: BytesIO,
        parse_images: bool = True,
    ) -> ParseResult:
        doc = Document(file_content)

        lines: List[Line] = []
        images: List[ImageArtefact] = []
        line_no = 0

        # 1. Основной текст (заголовки/абзацы)
        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style else ""
            if m := _HEADING_RE.match(style):
                level = int(m.group(1))
                md_prefix = "#" * level
                block_type = f"h{level}"
            else:
                md_prefix = ""
                block_type = "paragraph"

            text = para.text.strip()
            if not text and not md_prefix:
                continue

            content = f"{md_prefix} {text}".strip()
            lines.append(
                Line(
                    line_no=line_no,
                    block_type=block_type,
                    content=content,
                )
            )
            line_no += 1

        # 2. Таблицы
        for table in doc.tables:
            header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
            sep = " | ".join("---" for _ in table.rows[0].cells)
            lines.append(Line(line_no=line_no, block_type="table", content=header))
            line_no += 1
            lines.append(Line(line_no=line_no, block_type="table", content=sep))
            line_no += 1
            for row in table.rows[1:]:
                row_txt = " | ".join(cell.text.strip() for cell in row.cells)
                lines.append(Line(line_no=line_no, block_type="table", content=row_txt))
                line_no += 1

        # 3. Изображения
        if parse_images:
            rels = doc.part._rels  # type: ignore[attr-defined]
            for rel in rels.values():
                if rel.reltype == RT.IMAGE:
                    image_bytes = rel.target_part.blob
                    key = f"{doc_id}/images/{uuid4().hex}.{rel.target_part.filename.split('.')[-1]}"
                    block_id = f"img/{uuid4().hex}"
                    images.append(
                        ImageArtefact(
                            key=key,
                            data=image_bytes,
                            source_block_id=block_id,
                        )
                    )
                    # Добавляем строку-заглушку в MD
                    md_stub = f"![]({key})"
                    lines.append(
                        Line(
                            line_no=line_no,
                            block_type="image",
                            content=md_stub,
                            block_id=block_id,
                        )
                    )
                    line_no += 1

        return ParseResult(lines=lines, images=images)